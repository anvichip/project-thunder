# template_extractor.py - Extract HTML/CSS from PDF/DOCX
import os
import subprocess
import tempfile
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
import pdfplumber
import re

def extract_html_css_from_pdf(pdf_path: str) -> tuple[str, str]:
    """
    Extract HTML and CSS from PDF file
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Tuple of (html_content, css_content)
    """
    print(f"📄 Extracting template from PDF: {pdf_path}")
    
    try:
        # Method 1: Using pdf2htmlEX (if available)
        html_content = extract_with_pdf2htmlex(pdf_path)
        if html_content:
            return separate_html_css(html_content)
    except Exception as e:
        print(f"⚠️ pdf2htmlEX failed: {e}")
    
    try:
        # Method 2: Fallback to pdfplumber for text extraction
        html_content = extract_with_pdfplumber(pdf_path)
        return separate_html_css(html_content)
    except Exception as e:
        print(f"❌ All PDF extraction methods failed: {e}")
        raise

def extract_with_pdf2htmlex(pdf_path: str) -> str:
    """Extract using pdf2htmlEX tool"""
    with tempfile.TemporaryDirectory() as temp_dir:
        output_name = "output"
        
        # Run pdf2htmlEX
        result = subprocess.run(
            [
                'pdf2htmlEX',
                '--zoom', '1.3',
                '--dest-dir', temp_dir,
                pdf_path,
                output_name
            ],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode != 0:
            raise Exception(f"pdf2htmlEX failed: {result.stderr}")
        
        # Read generated HTML
        html_file = os.path.join(temp_dir, f"{output_name}.html")
        if not os.path.exists(html_file):
            raise Exception("HTML file not generated")
        
        with open(html_file, 'r', encoding='utf-8') as f:
            return f.read()

def extract_with_pdfplumber(pdf_path: str) -> str:
    """Extract text and create basic HTML using pdfplumber"""
    print("📄 Using pdfplumber for extraction...")
    
    html_parts = ['<!DOCTYPE html>\n<html>\n<head>\n<meta charset="UTF-8">\n</head>\n<body>\n']
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Extract text
            text = page.extract_text()
            if text:
                # Basic parsing
                lines = text.split('\n')
                for line in lines:
                    if line.strip():
                        # Detect headings (all caps or short lines)
                        if line.isupper() or len(line) < 50:
                            html_parts.append(f'<h2 class="section-header">{line.strip()}</h2>\n')
                        else:
                            html_parts.append(f'<p class="content">{line.strip()}</p>\n')
            
            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                html_parts.append('<table class="resume-table">\n')
                for row in table:
                    html_parts.append('<tr>\n')
                    for cell in row:
                        html_parts.append(f'<td>{cell or ""}</td>\n')
                    html_parts.append('</tr>\n')
                html_parts.append('</table>\n')
    
    html_parts.append('</body>\n</html>')
    return ''.join(html_parts)

def extract_html_css_from_docx(docx_path: str) -> tuple[str, str]:
    """
    Extract HTML and CSS from DOCX file
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Tuple of (html_content, css_content)
    """
    print(f"📝 Extracting template from DOCX: {docx_path}")
    
    doc = Document(docx_path)
    html_parts = []
    css_rules = []
    
    # Start HTML
    html_parts.append('<!DOCTYPE html>\n<html>\n<head>\n<meta charset="UTF-8">\n</head>\n<body>\n')
    
    # Process paragraphs
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        
        # Determine style
        style_class = "normal"
        if para.style.name.startswith('Heading'):
            level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
            style_class = f"heading-{level}"
            tag = f"h{level}"
        else:
            tag = "p"
        
        # Extract formatting
        if para.runs:
            run = para.runs[0]
            font_size = run.font.size.pt if run.font.size else 11
            font_name = run.font.name or 'Arial'
            is_bold = run.font.bold
            is_italic = run.font.italic
            color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            
            # Build CSS
            css = f".{style_class} {{"
            css += f"font-size: {font_size}pt; "
            css += f"font-family: '{font_name}'; "
            if is_bold:
                css += "font-weight: bold; "
            if is_italic:
                css += "font-style: italic; "
            if color:
                css += f"color: rgb({color[0]}, {color[1]}, {color[2]}); "
            css += "}"
            
            if css not in css_rules:
                css_rules.append(css)
        
        html_parts.append(f'<{tag} class="{style_class}">{para.text}</{tag}>\n')
    
    # Process tables
    for table in doc.tables:
        html_parts.append('<table class="resume-table">\n')
        for row in table.rows:
            html_parts.append('<tr>\n')
            for cell in row.cells:
                html_parts.append(f'<td>{cell.text}</td>\n')
            html_parts.append('</tr>\n')
        html_parts.append('</table>\n')
    
    html_parts.append('</body>\n</html>')
    
    # Create CSS
    css_content = '\n'.join(css_rules)
    css_content += '''
/* Base styles */
body {
    font-family: Arial, sans-serif;
    line-height: 1.6;
    margin: 40px;
    color: #333;
}

.resume-table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
}

.resume-table td {
    padding: 8px;
    border: 1px solid #ddd;
}

.section-header {
    color: #2c3e50;
    border-bottom: 2px solid #3498db;
    padding-bottom: 5px;
    margin-top: 20px;
}

.content {
    margin: 10px 0;
}
'''
    
    html_content = ''.join(html_parts)
    
    return html_content, css_content

def separate_html_css(full_html: str) -> tuple[str, str]:
    """
    Separate HTML and CSS from full HTML document
    
    Args:
        full_html: Complete HTML document
        
    Returns:
        Tuple of (html_template, css_template)
    """
    soup = BeautifulSoup(full_html, 'html.parser')
    
    # Extract CSS
    css_parts = []
    
    # Extract inline styles from style tags
    for style_tag in soup.find_all('style'):
        css_parts.append(style_tag.string or '')
        style_tag.decompose()  # Remove from HTML
    
    # Extract linked stylesheets content (if any)
    for link_tag in soup.find_all('link', {'rel': 'stylesheet'}):
        link_tag.decompose()
    
    # Get remaining HTML
    html_template = str(soup)
    
    # Add placeholders for dynamic content
    html_template = add_placeholders(html_template)
    
    css_template = '\n'.join(css_parts)
    
    # Add default CSS if empty
    if not css_template.strip():
        css_template = get_default_css()
    
    return html_template, css_template

def add_placeholders(html_content: str) -> str:
    """
    Add placeholders to HTML for dynamic content
    
    Args:
        html_content: HTML string
        
    Returns:
        HTML with placeholders
    """
    # This is a smart replacement that identifies likely dynamic fields
    # You can customize this based on common resume patterns
    
    # Common name patterns
    html_content = re.sub(
        r'<h1[^>]*>([A-Z][a-z]+ [A-Z][a-z]+)</h1>',
        '<h1 class="name">{{NAME}}</h1>',
        html_content,
        count=1
    )
    
    # Email patterns
    html_content = re.sub(
        r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
        '{{EMAIL}}',
        html_content
    )
    
    # Phone patterns
    html_content = re.sub(
        r'(\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9})',
        '{{PHONE}}',
        html_content
    )
    
    return html_content

def get_default_css() -> str:
    """Return default CSS for resumes"""
    return '''
/* Default Resume CSS */
body {
    font-family: 'Arial', 'Helvetica', sans-serif;
    line-height: 1.6;
    color: #333;
    max-width: 800px;
    margin: 0 auto;
    padding: 40px 20px;
    background: #fff;
}

.name {
    font-size: 32px;
    font-weight: bold;
    color: #2c3e50;
    margin-bottom: 10px;
    text-align: center;
}

.contact-info {
    text-align: center;
    margin-bottom: 30px;
    font-size: 14px;
    color: #666;
}

.section-header {
    font-size: 20px;
    font-weight: bold;
    color: #2c3e50;
    border-bottom: 2px solid #3498db;
    padding-bottom: 5px;
    margin-top: 30px;
    margin-bottom: 15px;
}

.content {
    margin: 10px 0;
    font-size: 14px;
}

.resume-table {
    width: 100%;
    border-collapse: collapse;
    margin: 15px 0;
}

.resume-table td {
    padding: 8px;
    border: 1px solid #ddd;
    font-size: 14px;
}

ul, ol {
    margin: 10px 0;
    padding-left: 25px;
}

li {
    margin: 5px 0;
}

@media print {
    body {
        margin: 0;
        padding: 20px;
    }
}
'''