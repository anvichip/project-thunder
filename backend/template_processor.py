# template_processor.py
import os
import re
from datetime import datetime
from docx import Document
import PyPDF2
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def process_template(template_path: str, profile_data: dict, output_format: str = 'pdf') -> str:
    """
    Process template and fill with profile data
    
    Args:
        template_path: Path to template file (PDF or DOCX)
        profile_data: Dictionary containing user profile data
        output_format: Output format ('pdf' or 'docx')
        
    Returns:
        Path to generated resume file
    """
    file_ext = os.path.splitext(template_path)[1].lower()
    
    if file_ext == '.docx':
        return process_docx_template(template_path, profile_data, output_format)
    elif file_ext == '.pdf':
        return process_pdf_template(template_path, profile_data, output_format)
    else:
        raise ValueError(f"Unsupported template format: {file_ext}")


def process_docx_template(template_path: str, profile_data: dict, output_format: str) -> str:
    """
    Process DOCX template
    """
    doc = Document(template_path)
    
    # Create placeholder mapping
    placeholders = {
        '{{NAME}}': profile_data.get('fullName', ''),
        '{{EMAIL}}': profile_data.get('email', ''),
        '{{PHONE}}': profile_data.get('phone', ''),
        '{{LINKEDIN}}': profile_data.get('linkedin', ''),
        '{{GITHUB}}': profile_data.get('github', ''),
        '{{SKILLS}}': profile_data.get('skills', ''),
        '{{EXPERIENCE}}': profile_data.get('experience', ''),
        '{{EDUCATION}}': profile_data.get('education', ''),
    }
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    
    # Save output
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"resume_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    
    # Convert to PDF if needed
    if output_format == 'pdf':
        pdf_path = convert_docx_to_pdf(output_path)
        os.remove(output_path)
        return pdf_path
    
    return output_path


def process_pdf_template(template_path: str, profile_data: dict, output_format: str) -> str:
    """
    Process PDF template (creates overlay)
    Note: PDF editing is limited. For best results, use DOCX templates.
    """
    # Create new PDF with profile data
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"resume_{timestamp}.pdf"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    
    # Add profile data to PDF
    y_position = height - 100
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, y_position, profile_data.get('fullName', ''))
    
    y_position -= 30
    c.setFont("Helvetica", 10)
    c.drawString(100, y_position, f"Email: {profile_data.get('email', '')}")
    
    y_position -= 20
    c.drawString(100, y_position, f"Phone: {profile_data.get('phone', '')}")
    
    if profile_data.get('linkedin'):
        y_position -= 20
        c.drawString(100, y_position, f"LinkedIn: {profile_data.get('linkedin', '')}")
    
    if profile_data.get('github'):
        y_position -= 20
        c.drawString(100, y_position, f"GitHub: {profile_data.get('github', '')}")
    
    # Skills
    if profile_data.get('skills'):
        y_position -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_position, "Skills")
        y_position -= 20
        c.setFont("Helvetica", 10)
        skills_lines = simpleSplit(profile_data.get('skills', ''), "Helvetica", 10, width - 200)
        for line in skills_lines[:3]:
            c.drawString(100, y_position, line)
            y_position -= 15
    
    # Experience
    if profile_data.get('experience'):
        y_position -= 30
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_position, "Experience")
        y_position -= 20
        c.setFont("Helvetica", 10)
        exp_lines = simpleSplit(profile_data.get('experience', ''), "Helvetica", 10, width - 200)
        for line in exp_lines[:10]:
            c.drawString(100, y_position, line)
            y_position -= 15
            if y_position < 100:
                break
    
    # Education
    if profile_data.get('education') and y_position > 100:
        y_position -= 30
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_position, "Education")
        y_position -= 20
        c.setFont("Helvetica", 10)
        edu_lines = simpleSplit(profile_data.get('education', ''), "Helvetica", 10, width - 200)
        for line in edu_lines[:5]:
            c.drawString(100, y_position, line)
            y_position -= 15
            if y_position < 100:
                break
    
    c.save()
    return output_path


def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert DOCX to PDF
    Note: This is a simplified version. In production, use LibreOffice or similar.
    """
    try:
        import subprocess
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # Try using LibreOffice (if available)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path],
            capture_output=True,
            timeout=30
        )
        
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            raise Exception("PDF conversion failed")
            
    except Exception as e:
        print(f"LibreOffice conversion failed: {e}")
        # Fallback: return docx if conversion fails
        return docx_path


def extract_placeholders(template_path: str) -> list:
    """
    Extract placeholders from template
    
    Args:
        template_path: Path to template file
        
    Returns:
        List of placeholder names
    """
    file_ext = os.path.splitext(template_path)[1].lower()
    placeholders = set()
    
    if file_ext == '.docx':
        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
            placeholders.update(matches)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                    placeholders.update(matches)
    
    return list(placeholders)


def validate_template(template_path: str) -> tuple[bool, str]:
    """
    Validate template file
    
    Args:
        template_path: Path to template file
        
    Returns:
        Tuple of (is_valid, message)
    """
    if not os.path.exists(template_path):
        return False, "Template file not found"
    
    file_ext = os.path.splitext(template_path)[1].lower()
    if file_ext not in ['.pdf', '.docx', '.doc']:
        return False, "Unsupported template format. Use PDF or DOCX."
    
    try:
        placeholders = extract_placeholders(template_path)
        return True, f"Valid template with {len(placeholders)} placeholders"
    except Exception as e:
        return False, f"Error reading template: {str(e)}"