# template_extractor_smart.py - Extract HTML preserving EXACT formatting
import os
import pdfplumber
from docx import Document
from parser.llm_client import call_ollama
import json
import base64
from io import BytesIO

def extract_html_from_resume(file_path: str) -> str:
    """
    Extract HTML representation from resume using LLM to preserve exact formatting
    
    Args:
        file_path: Path to PDF or DOCX file
        
    Returns:
        HTML string preserving original design exactly
    """
    print(f"🎨 Extracting HTML template from: {file_path}")
    
    # Extract text with detailed formatting info
    if file_path.lower().endswith('.pdf'):
        text_content, layout_info = extract_detailed_from_pdf(file_path)
    else:
        text_content, layout_info = extract_detailed_from_docx(file_path)
    
    # Use LLM to convert to HTML preserving EXACT structure
    html_content = convert_to_html_with_llm(text_content, layout_info)
    
    return html_content

def extract_detailed_from_pdf(pdf_path: str) -> tuple:
    """Extract text with detailed layout information from PDF"""
    print("📄 Extracting detailed layout from PDF...")
    
    content_parts = []
    layout_info = {
        "pages": [],
        "fonts": set(),
        "colors": set(),
        "alignments": []
    }
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_info = {
                "page_num": page_num + 1,
                "width": page.width,
                "height": page.height,
                "elements": []
            }
            
            # Extract text with positioning
            text = page.extract_text()
            if text:
                content_parts.append(f"[PAGE {page_num + 1}]\n{text}\n")
            
            # Extract words with positioning and styling
            words = page.extract_words(keep_blank_chars=True)
            for word in words:
                element = {
                    "text": word.get('text', ''),
                    "x0": word.get('x0', 0),
                    "y0": word.get('y0', 0),
                    "x1": word.get('x1', 0),
                    "y1": word.get('y1', 0),
                    "fontname": word.get('fontname', ''),
                    "size": word.get('size', 0)
                }
                page_info["elements"].append(element)
                
                if word.get('fontname'):
                    layout_info["fonts"].add(word['fontname'])
            
            layout_info["pages"].append(page_info)
    
    layout_info["fonts"] = list(layout_info["fonts"])
    
    return '\n'.join(content_parts), layout_info

def extract_detailed_from_docx(docx_path: str) -> tuple:
    """Extract text with detailed formatting from DOCX"""
    print("📝 Extracting detailed layout from DOCX...")
    
    doc = Document(docx_path)
    content_parts = []
    layout_info = {
        "fonts": set(),
        "colors": set(),
        "styles": [],
        "alignments": []
    }
    
    for para in doc.paragraphs:
        if para.text.strip():
            # Determine style and formatting
            style_info = {
                "text": para.text,
                "style": para.style.name,
                "alignment": str(para.alignment) if para.alignment else "LEFT"
            }
            
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                content_parts.append(f"[HEADING{level}] {para.text}")
            else:
                content_parts.append(para.text)
            
            # Extract run-level formatting
            for run in para.runs:
                if run.font.name:
                    layout_info["fonts"].add(run.font.name)
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    layout_info["colors"].add(f"rgb({rgb[0]},{rgb[1]},{rgb[2]})")
            
            layout_info["styles"].append(style_info)
    
    # Extract tables
    for table in doc.tables:
        content_parts.append("[TABLE]")
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            content_parts.append(row_text)
        content_parts.append("[/TABLE]")
    
    layout_info["fonts"] = list(layout_info["fonts"])
    layout_info["colors"] = list(layout_info["colors"])
    
    return '\n'.join(content_parts), layout_info

def convert_to_html_with_llm(text_content: str, layout_info: dict) -> str:
    """Use LLM to convert resume text to HTML preserving EXACT structure"""
    print("🤖 Converting to HTML using LLM with layout preservation...")
    
    # Create detailed layout description
    layout_desc = create_layout_description(layout_info)
    
    prompt = f"""You are an expert at converting resume content to HTML while preserving the EXACT original visual layout, formatting, and design.

ORIGINAL RESUME CONTENT:
{text_content}

LAYOUT INFORMATION:
{layout_desc}

YOUR TASK:
Create a complete HTML document that EXACTLY replicates the original resume's visual appearance.

CRITICAL REQUIREMENTS:
1. Preserve EXACT spacing, margins, padding, and alignment
2. Match font sizes, weights, and families precisely
3. Maintain color schemes exactly
4. Keep the same visual hierarchy (headings, subheadings, body text)
5. Replicate any special formatting (bold, italic, underline)
6. Use semantic HTML5 tags appropriately
7. Include comprehensive inline CSS to match original styling
8. Make it print-friendly and professional

STRUCTURE:
- Use <div class="resume-container"> as main wrapper
- Use <div class="resume-header"> for name and contact info
- Use <div class="resume-section"> for each major section
- Use appropriate heading tags (h1 for name, h2 for sections, h3 for subsections)
- Include all CSS inline in <style> tag
- Use placeholders like {{{{NAME}}}}, {{{{EMAIL}}}}, {{{{PHONE}}}} ONLY for contact info that will be dynamically filled

CSS REQUIREMENTS:
- Set exact font families, sizes, and weights
- Define precise margins and padding
- Match colors exactly
- Include print media queries
- Make responsive for different screen sizes

Return ONLY the complete HTML document starting with <!DOCTYPE html>. No explanations or markdown formatting."""

    try:
        html_response = call_ollama(prompt)
        
        # Clean up response
        html_response = html_response.strip()
        if html_response.startswith('```html'):
            html_response = html_response[7:]
        if html_response.startswith('```'):
            html_response = html_response[3:]
        if html_response.endswith('```'):
            html_response = html_response[:-3]
        
        html_response = html_response.strip()
        
        # Ensure it starts with DOCTYPE
        if not html_response.startswith('<!DOCTYPE'):
            html_response = '<!DOCTYPE html>\n' + html_response
        
        print(f"✅ HTML template generated ({len(html_response)} chars)")
        return html_response
        
    except Exception as e:
        print(f"❌ LLM conversion failed: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to enhanced basic HTML
        return create_enhanced_html_fallback(text_content, layout_info)

def create_layout_description(layout_info: dict) -> str:
    """Create a text description of layout information"""
    desc_parts = []
    
    if "fonts" in layout_info and layout_info["fonts"]:
        desc_parts.append(f"Fonts used: {', '.join(layout_info['fonts'])}")
    
    if "colors" in layout_info and layout_info["colors"]:
        desc_parts.append(f"Colors used: {', '.join(layout_info['colors'])}")
    
    if "pages" in layout_info:
        desc_parts.append(f"Number of pages: {len(layout_info['pages'])}")
        if layout_info["pages"]:
            page = layout_info["pages"][0]
            desc_parts.append(f"Page dimensions: {page.get('width', 0)}x{page.get('height', 0)}")
    
    if "styles" in layout_info and layout_info["styles"]:
        unique_styles = set(s.get('style', '') for s in layout_info["styles"])
        desc_parts.append(f"Document styles: {', '.join(unique_styles)}")
    
    return '\n'.join(desc_parts)

def create_enhanced_html_fallback(text_content: str, layout_info: dict) -> str:
    """Create enhanced HTML if LLM fails"""
    print("⚠️ Using enhanced fallback HTML generation...")
    
    lines = text_content.split('\n')
    
    # Detect fonts
    primary_font = layout_info.get("fonts", ["Arial"])[0] if layout_info.get("fonts") else "Arial"
    
    html_parts = [
        '<!DOCTYPE html>',
        '<html>',
        '<head>',
        '<meta charset="UTF-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        '<style>',
        '* { margin: 0; padding: 0; box-sizing: border-box; }',
        f'body {{ font-family: "{primary_font}", Arial, sans-serif; max-width: 850px; margin: 0 auto; padding: 40px 60px; line-height: 1.6; color: #333; background: #fff; }}',
        '.resume-container { background: white; }',
        '.resume-header { text-align: center; margin-bottom: 30px; border-bottom: 2px solid #333; padding-bottom: 20px; }',
        '.name { font-size: 32px; font-weight: bold; margin-bottom: 10px; color: #1a1a1a; letter-spacing: 1px; }',
        '.contact-info { font-size: 14px; color: #555; margin-top: 10px; }',
        '.contact-info a { color: #0066cc; text-decoration: none; }',
        '.resume-section { margin: 25px 0; }',
        '.section-header { font-size: 20px; font-weight: bold; border-bottom: 2px solid #666; margin: 25px 0 15px 0; padding-bottom: 5px; color: #2c3e50; text-transform: uppercase; letter-spacing: 0.5px; }',
        '.subsection { margin: 15px 0 15px 20px; }',
        '.subsection-title { font-size: 16px; font-weight: bold; color: #34495e; margin-bottom: 8px; }',
        '.subsection-content { margin-left: 20px; }',
        '.subsection-content p { margin: 6px 0; }',
        '.subsection-content ul { margin: 8px 0; padding-left: 20px; }',
        '.subsection-content li { margin: 4px 0; }',
        'table { width: 100%; border-collapse: collapse; margin: 15px 0; }',
        'table td { padding: 8px; border: 1px solid #ddd; }',
        '@media print {',
        '  body { padding: 20px; }',
        '  .resume-section { page-break-inside: avoid; }',
        '}',
        '@media (max-width: 768px) {',
        '  body { padding: 20px; }',
        '  .name { font-size: 24px; }',
        '  .section-header { font-size: 18px; }',
        '}',
        '</style>',
        '</head>',
        '<body>',
        '<div class="resume-container">'
    ]
    
    # Parse content
    in_header = True
    in_table = False
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Handle table markers
        if line == '[TABLE]':
            in_table = True
            html_parts.append('<table>')
            continue
        elif line == '[/TABLE]':
            in_table = False
            html_parts.append('</table>')
            continue
        
        if in_table:
            # Parse table row
            cells = line.split('|')
            html_parts.append('<tr>')
            for cell in cells:
                html_parts.append(f'<td>{cell.strip()}</td>')
            html_parts.append('</tr>')
            continue
        
        # Handle headings
        if line.startswith('[HEADING'):
            level = line[8] if len(line) > 8 and line[8].isdigit() else '2'
            content = line[line.index(']')+1:].strip()
            if in_header:
                html_parts.append(f'<div class="resume-header"><h1 class="name">{content}</h1></div>')
                in_header = False
            else:
                if current_section:
                    html_parts.append('</div>')  # Close previous section
                html_parts.append('<div class="resume-section">')
                html_parts.append(f'<h2 class="section-header">{content}</h2>')
                current_section = content
            continue
        
        # Handle page markers
        if line.startswith('[PAGE'):
            continue
        
        # Handle header content
        if in_header:
            if '@' in line or any(char.isdigit() for char in line):
                html_parts.append(f'<div class="resume-header">')
                html_parts.append(f'<h1 class="name">{{{{NAME}}}}</h1>')
                html_parts.append(f'<div class="contact-info">{line}</div>')
                html_parts.append('</div>')
                in_header = False
            else:
                html_parts.append(f'<div class="resume-header">')
                html_parts.append(f'<h1 class="name">{line}</h1>')
            continue
        
        # Regular content
        if line.isupper() and len(line) < 50:
            # Likely a section header
            if current_section:
                html_parts.append('</div>')
            html_parts.append('<div class="resume-section">')
            html_parts.append(f'<h2 class="section-header">{line}</h2>')
            current_section = line
        else:
            # Regular paragraph
            html_parts.append(f'<p>{line}</p>')
    
    # Close any open sections
    if current_section:
        html_parts.append('</div>')
    
    html_parts.extend([
        '</div>',
        '</body>',
        '</html>'
    ])
    
    return '\n'.join(html_parts)