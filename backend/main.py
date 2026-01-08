# main.py - FIXED CORS and Error Handling
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import OAuth2PasswordBearer
from pydantic import BaseModel, EmailStr
from typing import Optional, List
from motor.motor_asyncio import AsyncIOMotorClient
from datetime import datetime, timedelta
from jose import JWTError, jwt
from passlib.context import CryptContext
import os
import traceback
from dotenv import load_dotenv
from bson import ObjectId
from typing import Dict, Any

# Import custom modules
from resume_parser import parse_resume

load_dotenv()

app = FastAPI(title="CareerHub API")

# CORS Configuration - FIXED
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://localhost:5173",
        "http://localhost:5174",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:5173",
        "http://127.0.0.1:5174"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# MongoDB Configuration
MONGODB_URL = os.getenv("MONGODB_URL", "mongodb://localhost:27017")
DATABASE_NAME = os.getenv("DATABASE_NAME", "careerhub")

print(f"Connecting to MongoDB: {MONGODB_URL}")
print(f"Database: {DATABASE_NAME}")

client = AsyncIOMotorClient(MONGODB_URL)
db = client[DATABASE_NAME]
users_collection = db.users
profiles_collection = db.profiles
templates_collection = db.templates
drafts_collection = db.drafts

# Security Configuration
SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-change-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 10080

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# Directories
UPLOAD_DIR = "uploads"
TEMP_DIR = "temp"
OUTPUT_DIR = "output"

for directory in [UPLOAD_DIR, TEMP_DIR, OUTPUT_DIR]:
    os.makedirs(directory, exist_ok=True)
    print(f"Created directory: {directory}")

# Pydantic Models
class UserRegister(BaseModel):
    username: str
    email: EmailStr
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class FirebaseAuthRequest(BaseModel):
    uid: str
    email: EmailStr
    name: str
    photoURL: Optional[str] = None

class Token(BaseModel):
    access_token: str
    token_type: str
    user: dict

class ProfileData(BaseModel):
    fullName: str
    email: EmailStr
    phone: str
    linkedin: Optional[str] = None
    github: Optional[str] = None
    skills: Optional[str] = None
    experience: Optional[str] = None
    education: Optional[str] = None

class SaveProfileRequest(BaseModel):
    email: EmailStr
    profileData: ProfileData
    selectedRoles: List[str]

# Helper Functions
def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    user = await users_collection.find_one({"email": email})
    if user is None:
        raise credentials_exception
    return user

# Startup and Shutdown
@app.on_event("startup")
async def startup_db_client():
    try:
        await client.admin.command('ping')
        print("✅ Connected to MongoDB successfully")
        
        # Create indexes
        await users_collection.create_index("email", unique=True)
        await profiles_collection.create_index("email", unique=True)
        await templates_collection.create_index("email")
        await drafts_collection.create_index("email")
        print("✅ Database indexes created")
    except Exception as e:
        print(f"❌ MongoDB connection failed: {e}")
        traceback.print_exc()

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
    print("MongoDB connection closed")

# Global exception handler
@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    print(f"Global exception: {exc}")
    traceback.print_exc()
    return JSONResponse(
        status_code=500,
        content={"detail": str(exc)}
    )

# Root endpoint
@app.get("/")
async def root():
    return {
        "message": "CareerHub API is running",
        "status": "ok",
        "endpoints": {
            "auth": "/api/auth/*",
            "profiles": "/api/user-profile/*",
            "resume": "/api/upload-resume"
        }
    }

# Health check
@app.get("/health")
async def health_check():
    try:
        await client.admin.command('ping')
        return {"status": "healthy", "database": "connected"}
    except Exception as e:
        return {"status": "unhealthy", "database": "disconnected", "error": str(e)}

# Authentication Endpoints
@app.post("/api/auth/register", response_model=Token)
async def register(user_data: UserRegister):
    try:
        print(f"Registration attempt for: {user_data.email}")
        
        # Check if user exists
        existing_user = await users_collection.find_one({"email": user_data.email})
        if existing_user:
            print(f"User already exists: {user_data.email}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Email already registered"
            )
        
        # Check username
        existing_username = await users_collection.find_one({"username": user_data.username})
        if existing_username:
            print(f"Username already taken: {user_data.username}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Username already taken"
            )
        
        # Create user
        hashed_password = get_password_hash(user_data.password)
        user_doc = {
            "username": user_data.username,
            "email": user_data.email,
            "password": hashed_password,
            "auth_provider": "email",
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await users_collection.insert_one(user_doc)
        user_doc["_id"] = result.inserted_id
        
        print(f"User created successfully: {user_data.email}")
        
        # Create access token
        access_token = create_access_token(data={"sub": user_data.email})
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": str(user_doc["_id"]),
                "username": user_doc["username"],
                "email": user_doc["email"],
                "auth_provider": user_doc["auth_provider"]
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Registration error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Registration failed: {str(e)}")

@app.post("/api/auth/login", response_model=Token)
async def login(user_data: UserLogin):
    try:
        print(f"Login attempt for: {user_data.email}")
        
        # Find user
        user = await users_collection.find_one({"email": user_data.email})
        if not user:
            print(f"User not found: {user_data.email}")
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect email or password"
            )
        
        # Verify password
        if not verify_password(user_data.password, user["password"]):
            print(f"Invalid password for: {user_data.email}")
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect email or password"
            )
        
        print(f"Login successful for: {user_data.email}")
        
        # Create access token
        access_token = create_access_token(data={"sub": user["email"]})
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": str(user["_id"]),
                "username": user.get("username", ""),
                "email": user["email"],
                "name": user.get("name", ""),
                "auth_provider": user.get("auth_provider", "email"),
                "picture": user.get("picture")
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Login error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Login failed: {str(e)}")

@app.post("/api/auth/firebase", response_model=Token)
async def firebase_auth(auth_data: FirebaseAuthRequest):
    try:
        print(f"Firebase auth for: {auth_data.email}")
        
        email = auth_data.email
        user = await users_collection.find_one({"email": email})
        
        if not user:
            print(f"Creating new user via Firebase: {email}")
            # Create new user from Firebase data
            user_doc = {
                "email": email,
                "username": auth_data.name,
                "name": auth_data.name,
                "firebase_uid": auth_data.uid,
                "picture": auth_data.photoURL,
                "auth_provider": "google",
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow()
            }
            result = await users_collection.insert_one(user_doc)
            user_doc["_id"] = result.inserted_id
            user = user_doc
            print(f"✅ Created new user via Google: {email}")
        else:
            print(f"Updating existing user via Firebase: {email}")
            # Update existing user with latest info
            await users_collection.update_one(
                {"email": email},
                {"$set": {
                    "firebase_uid": auth_data.uid,
                    "picture": auth_data.photoURL,
                    "name": auth_data.name,
                    "updated_at": datetime.utcnow()
                }}
            )
            print(f"✅ Updated existing user via Google: {email}")
        
        # Create JWT access token
        access_token = create_access_token(data={"sub": email})
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": str(user["_id"]),
                "username": user.get("username", user.get("name", "")),
                "email": user["email"],
                "name": user.get("name", ""),
                "auth_provider": user.get("auth_provider", "google"),
                "picture": user.get("picture")
            }
        }
    except Exception as e:
        print(f"Firebase auth error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Firebase auth failed: {str(e)}")

# Resume Upload
@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...), userId: str = ""):
    temp_path = ""
    try:
        print(f"Resume upload for user: {userId}")
        
        if not file.filename.endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only PDF files are allowed")
        
        temp_path = os.path.join(TEMP_DIR, f"temp_{file.filename}")
        with open(temp_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        print(f"Parsing resume: {temp_path}")
        extracted_data = parse_resume(temp_path)
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        print("Resume parsed successfully")
        return {
            "message": "Resume parsed successfully",
            "extractedData": extracted_data
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Resume upload error: {e}")
        traceback.print_exc()
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=str(e))

# Profile endpoints
@app.post("/api/save-user-profile")
async def save_user_profile(request: SaveProfileRequest):
    try:
        print(f"Saving profile for: {request.email}")
        
        existing_profile = await profiles_collection.find_one({"email": request.email})
        
        profile_data = {
            "email": request.email,
            "profileData": request.profileData.dict(),
            "selectedRoles": request.selectedRoles,
            "updatedAt": datetime.utcnow()
        }
        
        if existing_profile:
            await profiles_collection.update_one(
                {"email": request.email},
                {"$set": profile_data}
            )
            message = "Profile updated successfully"
            print(f"Profile updated: {request.email}")
        else:
            profile_data["createdAt"] = datetime.utcnow()
            await profiles_collection.insert_one(profile_data)
            message = "Profile created successfully"
            print(f"Profile created: {request.email}")
        
        return {"message": message, "email": request.email}
    except Exception as e:
        print(f"Save profile error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/user-profile/{email}")
async def get_user_profile(email: str):
    try:
        print(f"Fetching profile for: {email}")
        
        profile = await profiles_collection.find_one({"email": email})
        
        if not profile:
            print(f"Profile not found: {email}")
            raise HTTPException(status_code=404, detail="Profile not found")
        
        profile["_id"] = str(profile["_id"])
        print(f"Profile found: {email}")
        return profile
    except HTTPException:
        raise
    except Exception as e:
        print(f"Get profile error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    
# Add these to main.py after existing endpoints

from fastapi.responses import StreamingResponse
import io

# Standard Resume Template Generator
@app.post("/api/generate-standard-resume")
async def generate_standard_resume(request: dict):
    """Generate resume using standard template"""
    try:
        email = request.get('email')
        format_type = request.get('format', 'pdf')
        
        print(f"Generating standard resume for: {email}, format: {format_type}")
        
        # Get user profile
        profile = await profiles_collection.find_one({"email": email})
        if not profile:
            raise HTTPException(status_code=404, detail="Profile not found")
        
        profile_data = profile.get('profileData', {})
        
        # Generate resume using standard template
        if format_type == 'pdf':
            output_path = generate_standard_pdf(profile_data, email)
            media_type = 'application/pdf'
        else:
            output_path = generate_standard_docx(profile_data, email)
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Save template info to database
        template_doc = {
            "email": email,
            "name": "Standard Template",
            "format": format_type,
            "type": "standard",
            "createdAt": datetime.utcnow()
        }
        await templates_collection.insert_one(template_doc)
        
        print(f"Standard resume generated successfully: {output_path}")
        
        # Return file
        return FileResponse(
            output_path,
            media_type=media_type,
            filename=f"resume.{format_type}"
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Standard resume generation error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/preview-standard-resume")
async def preview_standard_resume(request: dict):
    """Preview resume using standard template"""
    try:
        email = request.get('email')
        
        print(f"Previewing standard resume for: {email}")
        
        # Get user profile
        profile = await profiles_collection.find_one({"email": email})
        if not profile:
            raise HTTPException(status_code=404, detail="Profile not found")
        
        profile_data = profile.get('profileData', {})
        
        # Generate PDF preview
        output_path = generate_standard_pdf(profile_data, email)
        
        print(f"Standard resume preview generated: {output_path}")
        
        # Return file
        return FileResponse(
            output_path,
            media_type='application/pdf',
            filename="preview.pdf"
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Standard resume preview error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# Helper function to generate standard PDF
def generate_standard_pdf(profile_data: dict, email: str) -> str:
    """Generate PDF using standard template"""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"resume_{email.split('@')[0]}_{timestamp}.pdf"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    
    y = height - 0.75 * inch
    
    # Header - Name
    c.setFont("Helvetica-Bold", 24)
    name = profile_data.get('fullName', 'Your Name')
    c.drawString(0.75 * inch, y, name)
    
    y -= 0.4 * inch
    
    # Contact Information
    c.setFont("Helvetica", 10)
    contact_info = []
    if profile_data.get('email'):
        contact_info.append(f"✉ {profile_data['email']}")
    if profile_data.get('phone'):
        contact_info.append(f"☎ {profile_data['phone']}")
    if profile_data.get('linkedin'):
        contact_info.append(f"💼 LinkedIn")
    if profile_data.get('github'):
        contact_info.append(f"💻 GitHub")
    
    contact_line = " | ".join(contact_info)
    c.drawString(0.75 * inch, y, contact_line)
    
    y -= 0.5 * inch
    
    # Separator line
    c.setStrokeColor(HexColor('#2563eb'))
    c.setLineWidth(2)
    c.line(0.75 * inch, y, width - 0.75 * inch, y)
    
    y -= 0.4 * inch
    
    # Skills Section
    if profile_data.get('skills'):
        c.setFont("Helvetica-Bold", 14)
        c.drawString(0.75 * inch, y, "SKILLS")
        y -= 0.25 * inch
        
        c.setFont("Helvetica", 10)
        skills = profile_data['skills']
        # Wrap skills text
        skills_lines = []
        current_line = ""
        for skill in skills.split(','):
            skill = skill.strip()
            test_line = current_line + (", " if current_line else "") + skill
            if c.stringWidth(test_line, "Helvetica", 10) < (width - 1.75 * inch):
                current_line = test_line
            else:
                if current_line:
                    skills_lines.append(current_line)
                current_line = skill
        if current_line:
            skills_lines.append(current_line)
        
        for line in skills_lines[:3]:  # Max 3 lines
            c.drawString(0.75 * inch, y, line)
            y -= 0.2 * inch
        
        y -= 0.2 * inch
    
    # Experience Section
    if profile_data.get('experience'):
        c.setFont("Helvetica-Bold", 14)
        c.drawString(0.75 * inch, y, "EXPERIENCE")
        y -= 0.25 * inch
        
        c.setFont("Helvetica", 10)
        exp_lines = profile_data['experience'].split('\n')
        for line in exp_lines[:15]:  # Max 15 lines
            if y < 1.5 * inch:
                c.showPage()
                y = height - inch
            if line.strip():
                # Wrap long lines
                if c.stringWidth(line, "Helvetica", 10) > (width - 1.75 * inch):
                    words = line.split()
                    current_line = ""
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        if c.stringWidth(test_line, "Helvetica", 10) < (width - 1.75 * inch):
                            current_line = test_line
                        else:
                            if current_line:
                                c.drawString(0.75 * inch, y, current_line)
                                y -= 0.2 * inch
                            current_line = word
                    if current_line:
                        c.drawString(0.75 * inch, y, current_line)
                        y -= 0.2 * inch
                else:
                    c.drawString(0.75 * inch, y, line)
                    y -= 0.2 * inch
        
        y -= 0.2 * inch
    
    # Education Section
    if profile_data.get('education') and y > 2 * inch:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(0.75 * inch, y, "EDUCATION")
        y -= 0.25 * inch
        
        c.setFont("Helvetica", 10)
        edu_lines = profile_data['education'].split('\n')
        for line in edu_lines[:10]:  # Max 10 lines
            if y < 1.5 * inch:
                break
            if line.strip():
                # Wrap long lines
                if c.stringWidth(line, "Helvetica", 10) > (width - 1.75 * inch):
                    words = line.split()
                    current_line = ""
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        if c.stringWidth(test_line, "Helvetica", 10) < (width - 1.75 * inch):
                            current_line = test_line
                        else:
                            if current_line:
                                c.drawString(0.75 * inch, y, current_line)
                                y -= 0.2 * inch
                            current_line = word
                    if current_line:
                        c.drawString(0.75 * inch, y, current_line)
                        y -= 0.2 * inch
                else:
                    c.drawString(0.75 * inch, y, line)
                    y -= 0.2 * inch
    
    c.save()
    return output_path

# Helper function to generate standard DOCX
def generate_standard_docx(profile_data: dict, email: str) -> str:
    """Generate DOCX using standard template"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"resume_{email.split('@')[0]}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(profile_data.get('fullName', 'Your Name'))
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(37, 99, 235)
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Contact Information
    contact_para = doc.add_paragraph()
    contact_info = []
    if profile_data.get('email'):
        contact_info.append(profile_data['email'])
    if profile_data.get('phone'):
        contact_info.append(profile_data['phone'])
    if profile_data.get('linkedin'):
        contact_info.append('LinkedIn: ' + profile_data['linkedin'])
    if profile_data.get('github'):
        contact_info.append('GitHub: ' + profile_data['github'])
    
    contact_run = contact_para.add_run(' | '.join(contact_info))
    contact_run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # Skills Section
    if profile_data.get('skills'):
        skills_heading = doc.add_paragraph()
        skills_heading_run = skills_heading.add_run('SKILLS')
        skills_heading_run.font.size = Pt(14)
        skills_heading_run.font.bold = True
        skills_heading_run.font.color.rgb = RGBColor(37, 99, 235)
        
        skills_para = doc.add_paragraph(profile_data['skills'])
        skills_para.style = 'Normal'
        doc.add_paragraph()  # Spacer
    
    # Experience Section
    if profile_data.get('experience'):
        exp_heading = doc.add_paragraph()
        exp_heading_run = exp_heading.add_run('EXPERIENCE')
        exp_heading_run.font.size = Pt(14)
        exp_heading_run.font.bold = True
        exp_heading_run.font.color.rgb = RGBColor(37, 99, 235)
        
        exp_para = doc.add_paragraph(profile_data['experience'])
        exp_para.style = 'Normal'
        doc.add_paragraph()  # Spacer
    
    # Education Section
    if profile_data.get('education'):
        edu_heading = doc.add_paragraph()
        edu_heading_run = edu_heading.add_run('EDUCATION')
        edu_heading_run.font.size = Pt(14)
        edu_heading_run.font.bold = True
        edu_heading_run.font.color.rgb = RGBColor(37, 99, 235)
        
        edu_para = doc.add_paragraph(profile_data['education'])
        edu_para.style = 'Normal'
    
    doc.save(output_path)
    return output_path

# Fix the existing generate-resume endpoint
@app.post("/api/generate-resume")
async def generate_resume(
    template: UploadFile = File(...),
    email: str = "",
    format: str = "pdf"
):
    """Generate resume from uploaded template"""
    temp_path = ""
    try:
        print(f"Generating resume for: {email}, format: {format}")
        
        # Get user profile
        profile = await profiles_collection.find_one({"email": email})
        if not profile:
            raise HTTPException(status_code=404, detail="Profile not found")
        
        profile_data = profile.get('profileData', {})
        
        # Save template
        temp_path = os.path.join(TEMP_DIR, f"template_{email}_{template.filename}")
        with open(temp_path, "wb") as buffer:
            content = await template.read()
            buffer.write(content)
        
        print(f"Template saved: {temp_path}")
        
        # Process template with profile data
        from template_processor import process_template
        output_path = process_template(temp_path, profile_data, format)
        
        # Save template to database
        template_doc = {
            "email": email,
            "name": template.filename,
            "format": format,
            "type": "custom",
            "path": temp_path,
            "createdAt": datetime.utcnow()
        }
        await templates_collection.insert_one(template_doc)
        
        print(f"Resume generated: {output_path}")
        
        # Return file
        return FileResponse(
            output_path,
            media_type='application/pdf' if format == 'pdf' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"resume.{format}"
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Resume generation error: {e}")
        traceback.print_exc()
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/templates/{email}")
async def get_saved_templates(email: str):
    """Get all saved templates for a user"""
    try:
        print(f"Fetching templates for: {email}")
        templates = await templates_collection.find({"email": email}).to_list(100)
        
        result = []
        for template in templates:
            result.append({
                "id": str(template["_id"]),
                "_id": str(template["_id"]),
                "email": template.get("email"),
                "name": template.get("name", "Resume Template"),
                "format": template.get("format", "pdf"),
                "type": template.get("type", "custom"),
                "createdAt": template.get("createdAt").isoformat() if template.get("createdAt") else datetime.utcnow().isoformat()
            })
        
        print(f"Found {len(result)} templates")
        return result
    except Exception as e:
        print(f"Get templates error: {e}")
        traceback.print_exc()
        # Return empty list instead of error
        return []

@app.put("/api/user-profile/{email}")
async def update_user_profile(email: str, request: Dict[str, Any]):
    """Update user profile data"""
    try:
        print(f"Updating profile for: {email}")
        
        profile_data = request.get('profileData')
        if not profile_data:
            raise HTTPException(status_code=400, detail="Profile data is required")
        
        result = await profiles_collection.update_one(
            {"email": email},
            {
                "$set": {
                    "profileData": profile_data,
                    "updatedAt": datetime.utcnow()
                }
            }
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Profile not found")
        
        print(f"Profile updated successfully: {email}")
        return {"message": "Profile updated successfully", "email": email}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Update profile error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/user-profile/{email}/roles")
async def update_user_roles(email: str, request: Dict[str, Any]):
    """Update user selected roles"""
    try:
        print(f"Updating roles for: {email}")
        
        selected_roles = request.get('selectedRoles')
        if not selected_roles:
            raise HTTPException(status_code=400, detail="Selected roles are required")
        
        result = await profiles_collection.update_one(
            {"email": email},
            {
                "$set": {
                    "selectedRoles": selected_roles,
                    "updatedAt": datetime.utcnow()
                }
            }
        )
        
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Profile not found")
        
        print(f"Roles updated successfully: {email}")
        return {"message": "Roles updated successfully", "email": email}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Update roles error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    
resumes_collection = db.resumes
resume_links_collection = db.resume_links

import secrets
import string

def generate_short_code(length=8):
    """Generate a random short code for resume links"""
    chars = string.ascii_letters + string.digits
    return ''.join(secrets.choice(chars) for _ in range(length))


@app.get("/api/latex-templates/{email}")
async def get_latex_templates(email: str):
    """Get all LaTeX templates/resumes for a user"""
    try:
        print(f"Fetching LaTeX templates for: {email}")
        resumes = await resumes_collection.find({"email": email}).sort("createdAt", -1).to_list(100)
        
        result = []
        for resume in resumes:
            result.append({
                "id": str(resume["_id"]),
                "name": resume.get("name", "Untitled Resume"),
                "content": resume.get("content", ""),
                "pdfUrl": resume.get("pdfUrl"),
                "createdAt": resume.get("createdAt").isoformat() if resume.get("createdAt") else datetime.utcnow().isoformat(),
                "updatedAt": resume.get("updatedAt").isoformat() if resume.get("updatedAt") else datetime.utcnow().isoformat()
            })
        
        print(f"Found {len(result)} templates")
        return result
    except Exception as e:
        print(f"Get LaTeX templates error: {e}")
        traceback.print_exc()
        return []

@app.post("/api/save-latex-resume")
async def save_latex_resume(request: Dict[str, Any]):
    """Save a compiled LaTeX resume"""
    try:
        email = request.get('email')
        name = request.get('name')
        content = request.get('content')
        pdf_url = request.get('pdfUrl')
        
        if not email or not content:
            raise HTTPException(status_code=400, detail="Email and content are required")
        
        resume_doc = {
            "email": email,
            "name": name or f"Resume {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "content": content,
            "pdfUrl": pdf_url,
            "createdAt": datetime.utcnow(),
            "updatedAt": datetime.utcnow()
        }
        
        result = await resumes_collection.insert_one(resume_doc)
        
        print(f"LaTeX resume saved: {result.inserted_id}")
        return {
            "message": "Resume saved successfully",
            "id": str(result.inserted_id)
        }
    except Exception as e:
        print(f"Save LaTeX resume error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate-resume-link")
async def generate_resume_link(request: Dict[str, Any]):
    """Generate a shareable link for a resume"""
    try:
        resume_id = request.get('resumeId')
        email = request.get('email')
        
        if not resume_id or not email:
            raise HTTPException(status_code=400, detail="Resume ID and email are required")
        
        # Check if resume exists
        resume = await resumes_collection.find_one({"_id": ObjectId(resume_id), "email": email})
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Generate short code
        short_code = generate_short_code()
        
        # Check if code already exists
        while await resume_links_collection.find_one({"shortCode": short_code}):
            short_code = generate_short_code()
        
        # Create link document
        link_doc = {
            "resumeId": ObjectId(resume_id),
            "email": email,
            "shortCode": short_code,
            "createdAt": datetime.utcnow(),
            "views": 0
        }
        
        await resume_links_collection.insert_one(link_doc)
        
        # Generate full URL
        base_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        share_url = f"{base_url}/view/{short_code}"
        
        print(f"Generated resume link: {share_url}")
        return {
            "shortCode": short_code,
            "shareUrl": share_url
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Generate resume link error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/resume-link/{short_code}")
async def get_resume_by_link(short_code: str):
    """Get resume by short code (public endpoint)"""
    try:
        print(f"Fetching resume for short code: {short_code}")
        
        # Find link
        link = await resume_links_collection.find_one({"shortCode": short_code})
        if not link:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Increment view count
        await resume_links_collection.update_one(
            {"shortCode": short_code},
            {"$inc": {"views": 1}}
        )
        
        # Get resume
        resume = await resumes_collection.find_one({"_id": link["resumeId"]})
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        print(f"Resume found for short code: {short_code}")
        return {
            "name": resume.get("name", "Resume"),
            "pdfUrl": resume.get("pdfUrl"),
            "createdAt": resume.get("createdAt").isoformat() if resume.get("createdAt") else None
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Get resume by link error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/resume-details/{resume_id}")
async def get_resume_details(resume_id: str, email: str):
    """Get detailed information about a resume"""
    try:
        print(f"Fetching resume details: {resume_id}")
        
        resume = await resumes_collection.find_one({"_id": ObjectId(resume_id), "email": email})
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Get share link if exists
        link = await resume_links_collection.find_one({"resumeId": ObjectId(resume_id)})
        share_info = None
        if link:
            base_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
            share_info = {
                "shortCode": link["shortCode"],
                "shareUrl": f"{base_url}/view/{link['shortCode']}",
                "views": link.get("views", 0)
            }
        
        return {
            "id": str(resume["_id"]),
            "name": resume.get("name", "Resume"),
            "content": resume.get("content", ""),
            "pdfUrl": resume.get("pdfUrl"),
            "createdAt": resume.get("createdAt").isoformat() if resume.get("createdAt") else None,
            "updatedAt": resume.get("updatedAt").isoformat() if resume.get("updatedAt") else None,
            "shareInfo": share_info
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Get resume details error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/resume/{resume_id}")
async def delete_resume(resume_id: str, email: str):
    """Delete a resume"""
    try:
        print(f"Deleting resume: {resume_id}")
        
        result = await resumes_collection.delete_one({"_id": ObjectId(resume_id), "email": email})
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Delete associated links
        await resume_links_collection.delete_many({"resumeId": ObjectId(resume_id)})
        
        print(f"Resume deleted: {resume_id}")
        return {"message": "Resume deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Delete resume error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    print(f"Starting server on port {port}")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")